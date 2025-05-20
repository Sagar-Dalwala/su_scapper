import os
import re
import sys
import logging
import json
from pathlib import Path
import pandas as pd
from docx import Document

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("docx_extraction.log"),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# Define paths
DOCX_DIR = Path("output/docx")
COMBINED_DOCX_DIR = Path("output/combined_docx")
OUTPUT_DIR = Path("output/training_data")

# Create output directory if it doesn't exist
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Keywords to identify important sections in syllabi
SECTION_KEYWORDS = {
    'course_objectives': ['course objectives', 'objectives of the course', 'learning objectives'],
    'course_outcomes': ['course outcomes', 'learning outcomes', 'expected outcomes'],
    'assessment': ['assessment', 'evaluation', 'grading', 'marks distribution'],
    'textbooks': ['textbook', 'recommended books', 'reference books', 'essential reading'],
    'schedule': ['schedule', 'course calendar', 'week by week', 'lecture plan'],
    'prerequisites': ['prerequisites', 'pre-requisites', 'prior knowledge'],
    'unit_details': ['unit', 'module', 'section']
}

def extract_text_from_docx(docx_path):
    """Extract full text from a DOCX file."""
    try:
        doc = Document(docx_path)
        return '\n'.join(para.text for para in doc.paragraphs if para.text.strip())
    except Exception as e:
        logger.error(f"Error extracting text from {docx_path}: {str(e)}")
        return ""

def extract_structured_data(docx_path):
    """Extract structured information from a DOCX file."""
    try:
        doc = Document(docx_path)
        
        # Extract basic metadata
        program = docx_path.parent.name
        filename = docx_path.name
        
        # Initialize data structure
        data = {
            "program": program,
            "filename": filename,
            "title": "",
            "sections": {}
        }
        
        # Extract heading as title
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                data["title"] = para.text.strip()
                break
        
        # Extract content by section
        current_section = None
        section_content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # Check if this paragraph is a section header
            para_lower = text.lower()
            matched_section = None
            
            for section, keywords in SECTION_KEYWORDS.items():
                if any(keyword in para_lower for keyword in keywords):
                    matched_section = section
                    break
            
            # If we found a new section and have content from previous section
            if matched_section and current_section and section_content:
                data["sections"][current_section] = '\n'.join(section_content)
                section_content = []
                current_section = matched_section
            
            # If we found a new section for the first time
            elif matched_section and not current_section:
                current_section = matched_section
            
            # Add content to current section if we have one
            if current_section:
                section_content.append(text)
        
        # Add the last section if needed
        if current_section and section_content:
            data["sections"][current_section] = '\n'.join(section_content)
        
        return data
    except Exception as e:
        logger.error(f"Error extracting structured data from {docx_path}: {str(e)}")
        return {"program": program, "filename": filename, "title": "", "sections": {}}

def create_qa_pairs(structured_data):
    """Create question-answer pairs for model training."""
    qa_pairs = []
    
    # Basic information
    program = structured_data.get("program", "")
    title = structured_data.get("title", "")
    
    # Add general questions
    qa_pairs.append({
        "question": f"What is the title of the {program} syllabus?",
        "answer": title if title else "Unknown"
    })
    
    # Add section-specific questions
    sections = structured_data.get("sections", {})
    
    # Course objectives
    if "course_objectives" in sections:
        qa_pairs.append({
            "question": f"What are the objectives of the {title} course?",
            "answer": sections["course_objectives"]
        })
    
    # Course outcomes
    if "course_outcomes" in sections:
        qa_pairs.append({
            "question": f"What are the learning outcomes of the {title} course?",
            "answer": sections["course_outcomes"]
        })
    
    # Assessment
    if "assessment" in sections:
        qa_pairs.append({
            "question": f"How is the {title} course assessed?",
            "answer": sections["assessment"]
        })
    
    # Textbooks
    if "textbooks" in sections:
        qa_pairs.append({
            "question": f"What are the recommended textbooks for the {title} course?",
            "answer": sections["textbooks"]
        })
    
    # Prerequisites
    if "prerequisites" in sections:
        qa_pairs.append({
            "question": f"What are the prerequisites for the {title} course?",
            "answer": sections["prerequisites"]
        })
    
    # Unit details - create a pair for each unit if possible
    if "unit_details" in sections:
        unit_text = sections["unit_details"]
        
        # Try to identify individual units with a simple pattern match
        unit_pattern = r'Unit\s+(\d+)[:\s]+(.*?)(?=Unit\s+\d+|$)'
        units = re.findall(unit_pattern, unit_text, re.IGNORECASE | re.DOTALL)
        
        if units:
            for unit_num, unit_content in units:
                qa_pairs.append({
                    "question": f"What is covered in Unit {unit_num} of the {title} course?",
                    "answer": unit_content.strip()
                })
        else:
            # If we couldn't identify individual units, use the whole section
            qa_pairs.append({
                "question": f"What topics are covered in the {title} course?",
                "answer": unit_text
            })
    
    return qa_pairs

def process_all_docx_files():
    """Process all DOCX files in the organized directories."""
    # Find all DOCX files recursively
    all_docx_files = []
    for program_dir in DOCX_DIR.iterdir():
        if program_dir.is_dir():
            all_docx_files.extend(list(program_dir.glob("*.docx")))
    
    logger.info(f"Found {len(all_docx_files)} DOCX files to process")
    
    # Extract structured data from each file
    all_structured_data = []
    all_qa_pairs = []
    
    for docx_path in all_docx_files:
        logger.info(f"Processing {docx_path}")
        
        # Extract structured data
        structured_data = extract_structured_data(docx_path)
        all_structured_data.append(structured_data)
        
        # Create QA pairs
        qa_pairs = create_qa_pairs(structured_data)
        all_qa_pairs.extend(qa_pairs)
    
    # Save all structured data
    structured_data_path = OUTPUT_DIR / "structured_syllabus_data.json"
    with open(structured_data_path, 'w', encoding='utf-8') as f:
        json.dump(all_structured_data, f, indent=2)
    
    # Save all QA pairs
    qa_pairs_path = OUTPUT_DIR / "syllabus_qa_pairs.jsonl"
    with open(qa_pairs_path, 'w', encoding='utf-8') as f:
        for qa_pair in all_qa_pairs:
            f.write(json.dumps(qa_pair) + '\n')
    
    # Create a summary dataframe
    summary_data = []
    for data in all_structured_data:
        row = {
            "Program": data.get("program", ""),
            "Title": data.get("title", ""),
            "Sections Found": ', '.join(data.get("sections", {}).keys())
        }
        summary_data.append(row)
    
    summary_df = pd.DataFrame(summary_data)
    summary_path = OUTPUT_DIR / "extraction_summary.csv"
    summary_df.to_csv(summary_path, index=False)
    
    logger.info(f"Saved structured data to {structured_data_path}")
    logger.info(f"Saved QA pairs to {qa_pairs_path}")
    logger.info(f"Saved summary to {summary_path}")
    logger.info(f"Generated {len(all_qa_pairs)} QA pairs from {len(all_structured_data)} documents")

def main():
    """Extract data from DOCX files for AI training."""
    process_all_docx_files()

if __name__ == "__main__":
    main() 