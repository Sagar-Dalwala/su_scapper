import os
import sys
import re
import logging
from pathlib import Path
import pandas as pd
import PyPDF2
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("pdf_to_docx_conversion.log"),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# Define paths
PDF_DIR = Path("output/pdfs")
DOCX_OUTPUT_DIR = Path("output/docx")
COMBINED_OUTPUT_DIR = Path("output/combined_docx")

# Create output directories if they don't exist
DOCX_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
COMBINED_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def extract_program_semester(filename):
    """Extract program and semester information from filename."""
    filename = filename.lower()
    
    # Common programs
    program_patterns = {
        'bsc_it': r'bsc[_\s]*it|bsc.*information\s*technology',
        'bsc_cs': r'bsc[_\s]*cs|bsc.*computer\s*science',
        'bsc_mb': r'bsc[_\s]*mb|b\.sc\.\s*mb',
        'bsc_chemistry': r'bsc[_\s]*chem|b\.sc\.\s*chemistry',
        'msc_it': r'msc[_\s]*it|msc.*information\s*technology',
        'msc_cs': r'msc[_\s]*cs|msc.*computer\s*science',
        'msc_mb': r'm\.sc\.\s*mb',
        'msc_wmt': r'msc[_\s]*wmt',
        'msc_ac': r'msc[_\s]*ac',
        'msc_biotech': r'msc[_\s]*biotech|m\.sc\.\s*biotechnology',
        'msc_genetics': r'm\.sc\.\s*genetics',
        'msc_organic_chemistry': r'm\.sc\.\s*organic\s*chemistry',
        'msc_medical_bt': r'm\.sc\.\s*medical\s*bt',
        'msc_clinical_embryology': r'm-sc-\s*clinical\s*embryology|m-sc\.\s*clinical',
        'pgdmlt': r'pgdmlt'
    }
    
    # Find program
    program = 'unknown'
    for prog, pattern in program_patterns.items():
        if re.search(pattern, filename, re.IGNORECASE):
            program = prog
            break
    
    # Find semester
    semester_match = re.search(r'sem[_\s.-]*(\d+)|semester[_\s.-]*(\d+)|sem[_\s.-]*([iv]+)', filename, re.IGNORECASE)
    semester = None
    
    if semester_match:
        # Check which group matched
        for group in semester_match.groups():
            if group:
                semester = group
                break
                
        # Convert roman numerals if needed
        if semester and semester.lower() in ['i', 'ii', 'iii', 'iv', 'v', 'vi']:
            roman_to_int = {'i': '1', 'ii': '2', 'iii': '3', 'iv': '4', 'v': '5', 'vi': '6'}
            semester = roman_to_int[semester.lower()]
    
    if not semester:
        semester = 'unknown'
        
    return program, semester

def convert_pdf_to_docx(pdf_path, docx_path):
    """Convert a PDF file to DOCX format."""
    logger.info(f"Converting {pdf_path} to {docx_path}")
    
    try:
        # Create a new Word document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Read PDF
        with open(pdf_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Add a title based on the file name
            title = os.path.splitext(os.path.basename(pdf_path))[0]
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Extract text from each page and add to document
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                
                # Add page content with page number
                doc.add_paragraph(f"Page {page_num + 1}")
                doc.add_paragraph(text)
                doc.add_paragraph("") # Add space between pages
        
        # Save the document
        doc.save(docx_path)
        return True
    except Exception as e:
        logger.error(f"Error converting {pdf_path}: {str(e)}")
        return False

def combine_docx_by_program(program_docs, output_path):
    """Combine multiple DOCX files for a program into one master document."""
    master_doc = Document()
    
    # Add program title
    program_title = program_docs[0].parent.name.upper().replace('_', ' ')
    heading = master_doc.add_heading(f"{program_title} - Combined Syllabus", level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Sort by semester
    sorted_docs = sorted(program_docs, key=lambda x: x.stem)
    
    for doc_path in sorted_docs:
        logger.info(f"Adding {doc_path} to combined document")
        
        try:
            # Add separator and document title
            master_doc.add_paragraph("=" * 80)
            title = doc_path.stem
            master_doc.add_heading(title, level=2)
            
            # Append content from the individual document
            doc = Document(doc_path)
            for element in doc.element.body:
                master_doc.element.body.append(element)
                
        except Exception as e:
            logger.error(f"Error adding {doc_path} to combined document: {str(e)}")
    
    # Save the combined document
    master_doc.save(output_path)
    logger.info(f"Saved combined document to {output_path}")

def main():
    """Convert all PDFs to DOCX and organize by program."""
    # Get all PDF files
    pdf_files = list(PDF_DIR.glob("*.pdf"))
    logger.info(f"Found {len(pdf_files)} PDF files to convert")
    
    if not pdf_files:
        logger.error(f"No PDF files found in {PDF_DIR}")
        return
    
    # Dictionary to track docx files by program
    program_docx_files = {}
    
    # Process each PDF
    for pdf_path in pdf_files:
        # Extract program and semester from filename
        program, semester = extract_program_semester(pdf_path.name)
        
        # Create program directory
        program_dir = DOCX_OUTPUT_DIR / program
        program_dir.mkdir(exist_ok=True)
        
        # Create DOCX path with organized naming
        docx_filename = f"{program}_sem{semester}_{pdf_path.stem}.docx"
        docx_path = program_dir / docx_filename
        
        # Convert PDF to DOCX
        success = convert_pdf_to_docx(pdf_path, docx_path)
        
        if success:
            # Track for combined document
            if program not in program_docx_files:
                program_docx_files[program] = []
            program_docx_files[program].append(docx_path)
    
    # Create combined documents by program
    logger.info("Creating combined documents by program")
    for program, docx_list in program_docx_files.items():
        if docx_list:
            combined_path = COMBINED_OUTPUT_DIR / f"{program}_combined_syllabus.docx"
            combine_docx_by_program(docx_list, combined_path)
    
    logger.info("Conversion and organization complete!")
    
    # Create a summary report
    summary = {
        "Program": [],
        "Semester Documents": [],
        "Combined Document": []
    }
    
    for program, docs in program_docx_files.items():
        summary["Program"].append(program)
        summary["Semester Documents"].append(len(docs))
        summary["Combined Document"].append("Yes" if docs else "No")
    
    summary_df = pd.DataFrame(summary)
    summary_path = COMBINED_OUTPUT_DIR / "conversion_summary.csv"
    summary_df.to_csv(summary_path, index=False)
    logger.info(f"Summary report saved to {summary_path}")

if __name__ == "__main__":
    main() 